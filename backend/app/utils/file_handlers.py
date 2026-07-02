import os
import logging
from typing import List, Dict, Any, Optional
import pymupdf
from docx import Document as DocxDocument

logger = logging.getLogger(__name__)


class TextBlock:
    """
    Standard block representing extracted text chunk containing page references,
    content, and specific metadata.
    """
    def __init__(self, content: str, page_number: Optional[int] = None, metadata: Optional[Dict[str, Any]] = None):
      self.content = content
      self.page_number = page_number
      self.metadata = metadata or {}


class FileExtractor:
    """
    Utility handler that parses binary and text payloads based on file extensions.
    """
    @staticmethod
    def extract_pdf(filepath: str) -> List[TextBlock]:
        """Extracts text page-by-page from a PDF file using PyMuPDF."""
        blocks: List[TextBlock] = []
        try:
            doc = pymupdf.open(filepath)
            for page_num in range(len(doc)):
                page = doc[page_num]
                text = page.get_text("text", sort=True)
                if text.strip():
                    blocks.append(TextBlock(
                        content=text,
                        page_number=page_num + 1, # 1-based page number
                        metadata={"page": page_num + 1}
                    ))
            doc.close()
        except Exception as e:
            logger.error(f"Error extracting PDF from {filepath}: {e}")
            raise RuntimeError(f"Failed to parse PDF: {str(e)}")
        return blocks

    @staticmethod
    def extract_docx(filepath: str) -> List[TextBlock]:
        """Extracts text from paragraphs and tables in a Word document."""
        blocks: List[TextBlock] = []
        try:
            doc = DocxDocument(filepath)
            
            # Extract paragraphs
            paragraphs_text = []
            for p in doc.paragraphs:
                if p.text.strip():
                    paragraphs_text.append(p.text)
            
            # Extract tables
            table_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                    if row_text:
                        table_text.append(" | ".join(row_text))

            full_text = "\n".join(paragraphs_text + table_text)
            if full_text.strip():
                # Word doc is typically treated as a single page or split logically
                blocks.append(TextBlock(
                    content=full_text,
                    page_number=1,
                    metadata={"page": 1}
                ))
        except Exception as e:
            logger.error(f"Error extracting DOCX from {filepath}: {e}")
            raise RuntimeError(f"Failed to parse DOCX: {str(e)}")
        return blocks

    @staticmethod
    def extract_text(filepath: str) -> List[TextBlock]:
        """Extracts text from TXT or MD files."""
        try:
            with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                content = f.read()
            if content.strip():
                return [TextBlock(content=content, page_number=1, metadata={"page": 1})]
        except Exception as e:
            logger.error(f"Error reading text file {filepath}: {e}")
            raise RuntimeError(f"Failed to parse text file: {str(e)}")
        return []

    @classmethod
    def extract(cls, filepath: str, filename: str) -> List[TextBlock]:
        """Dispatches file extraction based on file extension."""
        ext = os.path.splitext(filename.lower())[1]
        if ext == ".pdf":
            return cls.extract_pdf(filepath)
        elif ext in [".docx", ".doc"]:
            return cls.extract_docx(filepath)
        elif ext in [".txt", ".md", ".json", ".yaml", ".yml", ".py", ".java", ".js", ".ts"]:
            return cls.extract_text(filepath)
        else:
            raise ValueError(f"Unsupported file extension: {ext}")
